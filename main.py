import pymysql
import sys
from PyQt6 import uic  # Импортируем uic
from PyQt6.QtWidgets import QApplication, QMainWindow, QHeaderView, QTableWidgetItem, QDialog, QMessageBox
from PyQt6.QtCore import QDate
from docx import Document


class Database:
    def __init__(self):
        self.connection = None
        self.connect()

    def connect(self):
        try:
            self.connection = pymysql.connect(
                host='127.0.0.1',
                port=3307,
                user='root',
                password='123456',
                database='applicant_tracking',
                charset='utf8',
                autocommit=True
            )
            print('бд подключена')
        except Exception as e:
            self.connection = None
            print('бд не подключена:', e)

    def select(self, table):
        with self.connection.cursor() as cursor:
            cursor.execute(f"SELECT * FROM `{table}`")
            result = cursor.fetchall()
            return result

    def get_applications_with_specialties(self):
        with self.connection.cursor() as cursor:
            sql = """
                SELECT 
                    a.application_id,
                    a.snils,
                    s.speciality_name,
                    a.submission_date,
                    a.benefit,
                    a.study_type
                FROM applications a
                JOIN specialties s 
                    ON a.speciality_code = s.speciality_code
            """
            cursor.execute(sql)
            return cursor.fetchall()

    def get_specialties_with_departaments(self):
        with self.connection.cursor() as cursor:
            sql = """
                SELECT 
                    s.speciality_code,
                    s.speciality_name,
                    d.department_name,
                    s.budget_places,
                    s.paid_places,
                    s.study_duration,
                    s.study_form
                FROM specialties s
                JOIN departaments d 
                    ON s.department_id = d.department_id
            """
            cursor.execute(sql)
            return cursor.fetchall()

    def filter(self, table, filters):
        filtered = {
            k: v for k, v in filters.items()
            if v is not None and 'Выберите' not in str(v)
        }

        if not filtered:
            return self.select(table)

        conditions = []

        for key, value in filtered.items():
            if value in [0, 1, '0', '1']:
                conditions.append(f"`{key}` = {int(value)}")
            else:
                conditions.append(f"`{key}` = %s")

        where = " AND ".join(conditions)

        sql_map = {
            'applications': f"""
                SELECT a.application_id, a.snils, s.speciality_name,
                       a.submission_date, a.benefit, a.study_type
                FROM applications a
                JOIN specialties s 
                    ON a.speciality_code = s.speciality_code
                WHERE {where}
            """,
            'specialties': f"""
                SELECT s.speciality_code, s.speciality_name,
                       d.department_name, s.budget_places,
                       s.paid_places, s.study_duration, s.study_form
                FROM specialties s
                JOIN departaments d 
                    ON s.department_id = d.department_id
                WHERE {where}
            """,
            'applicants': f"SELECT * FROM applicants WHERE {where}"
        }

        sql = sql_map.get(table)
        if not sql:
            return []

        with self.connection.cursor() as cursor:
            cursor.execute(sql, tuple(
                v for v in filtered.values()
                if v not in [0, 1, '0', '1']
            ))
            return cursor.fetchall()

    def insert_applicant(self, data):
        with self.connection.cursor() as cursor:
            sql = """
                INSERT INTO applicants 
                (snils, last_name, first_name, middle_name, birth_date,
                gender, phone, passport, medical_certificate, email, address, foto)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """
            cursor.execute(sql, tuple(data.values()))

    def insert_education(self, data):
        with self.connection.cursor() as cursor:
            sql = """
                INSERT INTO education_documents
                (snils, school, graduation_year, average_grade)
                VALUES (%s, %s, %s, %s)
            """
            cursor.execute(sql, tuple(data.values()))

    def insert_application(self, data):
        with self.connection.cursor() as cursor:
            sql = """
                INSERT INTO applications 
                (snils, speciality_code, submission_date, benefit, study_type)
                VALUES (%s, %s, %s, %s, %s)
            """
            cursor.execute(sql, (
                data['snils'],
                data['speciality_code'],
                data['submission_date'],
                None if data['benefit'] == 'Нет' else data['benefit'],
                data.get('study_type')
            ))

    def insert_specialty(self, data):
        with self.connection.cursor() as cursor:
            sql = """
                INSERT INTO specialties
                (speciality_code, speciality_name, department_id,
                budget_places, paid_places, study_duration, study_form)
                VALUES (%s, %s, %s, %s, %s, %s, %s)
            """
            cursor.execute(sql, (
                data['code'],
                data['name'],
                data['dep_id'],
                data['budget'],
                data['paid'],
                data['duration'],
                data['form']
            ))
        self.connection.commit()

    def insert_department(self, data):
        """Добавление отделения"""
        with self.connection.cursor() as cursor:
            sql = """
                INSERT INTO departaments
                (department_name, head_of_department)
                VALUES (%s, %s)
            """
            cursor.execute(sql, (
                data['name'],
                data['head']
            ))
        self.connection.commit()

    def update_applicant(self, data):
        with self.connection.cursor() as cursor:
            sql = """
                UPDATE applicants SET
                last_name = %s,
                first_name = %s,
                middle_name = %s,
                birth_date = %s,
                gender = %s,
                phone = %s,
                passport = %s,
                medical_certificate = %s,
                email = %s,
                address = %s,
                foto = %s
                WHERE snils = %s"""
            cursor.execute(sql, (
                data['last_name'], data['first_name'],
                data['middle_name'], data['birth_date'], data['gender'],
                data['phone'], data['passport'], data['medical_certificate'],
                data['email'], data['address'], data['foto'],
                data['snils']
            ))

    def update_education(self, data):
        with self.connection.cursor() as cursor:
            sql = """
                UPDATE education_documents SET
                school = %s,
                graduation_year = %s,
                average_grade = %s
                WHERE snils = %s"""
            cursor.execute(sql, (
                data['school'],
                data['graduation_year'],
                data['average_grade'],
                data['snils']
            ))

    def update_application(self, application_id, data):
        with self.connection.cursor() as cursor:
            sql = """
                UPDATE applications SET
                    snils = %s,
                    speciality_code = %s,
                    submission_date = %s,
                    benefit = %s,
                    study_type = %s
                WHERE application_id = %s
            """
            cursor.execute(sql, (
                data['snils'],
                data['speciality_code'],
                data['submission_date'],
                None if data['benefit'] == 'Нет' else data['benefit'],
                data.get('study_type'),
                application_id
            ))

    def update_specialty(self, old_code, data):
        with self.connection.cursor() as cursor:
            sql = """
                UPDATE specialties SET
                    speciality_code = %s,
                    speciality_name = %s,
                    department_id = %s,
                    budget_places = %s,
                    paid_places = %s,
                    study_duration = %s,
                    study_form = %s
                WHERE speciality_code = %s
            """
            cursor.execute(sql, (
                data['code'],
                data['name'],
                data['dep_id'],
                data['budget'],
                data['paid'],
                data['duration'],
                data['form'],
                old_code
            ))
        self.connection.commit()

    def update_department(self, department_id, data):
        """Обновление отделения"""
        with self.connection.cursor() as cursor:
            sql = """
                UPDATE departaments SET
                    department_name = %s,
                    head_of_department = %s
                WHERE department_id = %s
            """
            cursor.execute(sql, (
                data['name'],
                data['head'],
                department_id
            ))
        self.connection.commit()

    def delete_applicant(self, snils):
        with self.connection.cursor() as cursor:
            cursor.execute(
                "DELETE FROM applicants WHERE snils = %s",
                (snils,)
            )

    def delete_application(self, application_id):
        with self.connection.cursor() as cursor:
            cursor.execute(
                "DELETE FROM applications WHERE application_id = %s",
                (application_id,)
            )

    def delete_department(self, department_id):
        """Удаление отделения"""
        with self.connection.cursor() as cursor:
            cursor.execute(
                "DELETE FROM departaments WHERE department_id = %s",
                (department_id,)
            )

    def application_exists(self, snils, speciality_code, study_type):
        with self.connection.cursor() as cursor:
            cursor.execute("""
                SELECT 1 FROM applications
                WHERE snils = %s
                AND speciality_code = %s
                AND study_type = %s
                LIMIT 1
            """, (snils, speciality_code, study_type))
            return cursor.fetchone() is not None

    def login_user(self, login, password):
        with self.connection.cursor() as cursor:
            cursor.execute("""
                SELECT role FROM users
                WHERE login = %s AND password = %s
            """, (login, password))

            return cursor.fetchone()

    def get_competition_list(self):
        with self.connection.cursor() as cursor:
            cursor.execute("""
                SELECT 
                    a.snils,
                    ap.last_name,
                    ap.first_name,
                    ap.middle_name,
                    s.speciality_name,
                    a.study_type,
                    a.benefit,
                    a.submission_date,
                    ed.average_grade
                FROM applications a
                JOIN applicants ap ON a.snils = ap.snils
                JOIN specialties s ON a.speciality_code = s.speciality_code
                JOIN education_documents ed ON ed.snils = a.snils
                ORDER BY ed.average_grade DESC, s.speciality_name, a.submission_date
            """)
            return cursor.fetchall()

    def get_enrollment_data(self):
        with self.connection.cursor() as cursor:
            cursor.execute("""
                SELECT 
                    a.snils,
                    ap.last_name,
                    ap.first_name,
                    ap.middle_name,
                    s.speciality_name,
                    a.study_type,
                    a.benefit,
                    a.submission_date,
                    ed.average_grade,
                    s.budget_places,
                    s.paid_places
                FROM applications a
                JOIN applicants ap ON a.snils = ap.snils
                JOIN specialties s ON a.speciality_code = s.speciality_code
                JOIN education_documents ed ON ed.snils = a.snils
                ORDER BY 
                    CASE WHEN a.benefit IS NOT NULL THEN 0 ELSE 1 END,
                    ed.average_grade DESC
            """)
            return cursor.fetchall()

    def close(self):
        if self.connection:
            self.connection.close()


class LoginWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        uic.loadUi('./ui/login.ui', self)
        self.setup_ui()
        self.login_btn.clicked.connect(self.login)
        self.main_window = None

    def setup_ui(self):
        self.setWindowTitle("Авторизация")
        self.login_input.setFocus()

    def login(self):
        login = self.login_input.text().strip()
        password = self.password_input.text().strip()

        result = Database().login_user(login, password)

        if result:
            role = result[0]

            self.main_window = MyWidget(role=role)
            self.main_window.show()
            self.close()
        else:
            QMessageBox.warning(self, "Ошибка", "Неверный логин или пароль")


class ApplicantDialog(QDialog):
    def __init__(self, db, snils=None, parent=None):
        super().__init__(parent)
        uic.loadUi('./ui/applicant_dialog.ui', self)
        self.db = db
        self.snils = snils
        self.btnSave.clicked.connect(self.save)
        self.btnCancel.clicked.connect(self.reject)
        self.inputSnils.setEnabled(True)

        if self.snils:
            self.setWindowTitle("Редактирование абитуриента")
            self.titleLabel.setText("Редактирование абитуриента")
            self.inputSnils.setEnabled(False)
            self.fill_data()
        else:
            self.setWindowTitle("Добавление абитуриента")
            self.inputSnils.setEnabled(True)

    def fill_data(self):
        '''Заполняем поля данными из БД'''
        with self.db.connection.cursor() as cursor:
            cursor.execute(
                "SELECT * FROM `applicants` WHERE snils = %s", (self.snils,))
            row = cursor.fetchone()

        if not row:
            return

        self.inputSnils.setText(str(row[0]))
        self.inputLastName.setText(str(row[1]))
        self.inputFirstName.setText(str(row[2]))
        self.inputMiddleName.setText(str(row[3]) if row[3] else '')
        self.inputBirthDate.setDate(row[4])

        gender_index = self.inputGender.findText(str(row[5]) if row[5] else '')
        self.inputGender.setCurrentIndex(
            gender_index if gender_index >= 0 else 0)

        self.inputPhone.setText(str(row[6]))
        self.inputPassport.setText(str(row[7]))
        self.inputMedical.setChecked(bool(row[8]))
        self.inputEmail.setText(str(row[9]))
        self.inputAddress.setText(str(row[10]))
        self.inputFoto.setChecked(bool(row[11]))

        with self.db.connection.cursor() as cursor:
            cursor.execute(
                "SELECT * FROM `education_documents` WHERE snils = %s", (self.snils,))
            edu = cursor.fetchone()

        if edu:
            self.inputSchool.setText(str(edu[3]))
            self.inputGraduationYear.setValue(int(edu[2]))
            self.inputAverageGrade.setValue(float(edu[4]))

    def validate(self):
        if not self.inputSnils.text().strip():
            QMessageBox.warning(self, "Ошибка", "Введите СНИЛС")
            return False
        if not self.inputLastName.text().strip():
            QMessageBox.warning(self, "Ошибка", "Введите фамилию")
            return False
        if not self.inputFirstName.text().strip():
            QMessageBox.warning(self, "Ошибка", "Введите имя")
            return False
        if not self.inputPhone.text().strip():
            QMessageBox.warning(self, "Ошибка", "Введите телефон")
            return False
        if not self.inputPassport.text().strip():
            QMessageBox.warning(self, "Ошибка", "Введите паспортные данные")
            return False
        if not self.inputEmail.text().strip():
            QMessageBox.warning(self, "Ошибка", "Введите email")
            return False
        if not self.inputAddress.text().strip():
            QMessageBox.warning(self, "Ошибка", "Введите адрес")
            return False
        if not self.inputSchool.text().strip():
            QMessageBox.warning(self, "Ошибка", "Введите название школы")
            return False
        phone = self.inputPhone.text().strip()
        if not phone:
            QMessageBox.warning(self, "Ошибка", "Введите телефон")
            return False
        if not phone.startswith('+'):
            QMessageBox.warning(
                self, "Ошибка", "Телефон должен начинаться с +")
            return False
        if len(phone) != 12:
            QMessageBox.warning(
                self, "Ошибка", "Телефон должен содержать ровно 12 символов")
            return False
        if not phone[1:].isdigit():
            QMessageBox.warning(
                self, "Ошибка", "Телефон должен содержать только цифры после +")
            return False
        return True

    def save(self):
        if not self.validate():
            return
        try:
            if self.snils:
                # РЕДАКТИРОВАНИЕ
                self.db.update_applicant({
                    'snils': self.snils,
                    'last_name': self.inputLastName.text().strip(),
                    'first_name': self.inputFirstName.text().strip(),
                    'middle_name': self.inputMiddleName.text().strip() or None,
                    'birth_date': self.inputBirthDate.date().toString("yyyy-MM-dd"),
                    'gender': self.inputGender.currentText() if self.inputGender.currentIndex() != 0 else None,
                    'phone': self.inputPhone.text().strip(),
                    'passport': self.inputPassport.text().strip(),
                    'medical_certificate': 1 if self.inputMedical.isChecked() else 0,
                    'email': self.inputEmail.text().strip(),
                    'address': self.inputAddress.text().strip(),
                    'foto': 1 if self.inputFoto.isChecked() else 0
                })

                self.db.update_education({
                    'snils': self.snils,
                    'school': self.inputSchool.text().strip(),
                    'graduation_year': self.inputGraduationYear.value(),
                    'average_grade': self.inputAverageGrade.value()
                })
                QMessageBox.information(
                    self, "Успех", "Данные успешно обновлены!")
            else:
                # ДОБАВЛЕНИЕ
                self.db.insert_applicant({
                    'snils': self.inputSnils.text().strip(),
                    'last_name': self.inputLastName.text().strip(),
                    'first_name': self.inputFirstName.text().strip(),
                    'middle_name': self.inputMiddleName.text().strip() or None,
                    'birth_date': self.inputBirthDate.date().toString("yyyy-MM-dd"),
                    'gender': self.inputGender.currentText() if self.inputGender.currentIndex() != 0 else None,
                    'phone': self.inputPhone.text().strip(),
                    'passport': self.inputPassport.text().strip(),
                    'medical_certificate': 1 if self.inputMedical.isChecked() else 0,
                    'email': self.inputEmail.text().strip(),
                    'address': self.inputAddress.text().strip(),
                    'foto': 1 if self.inputFoto.isChecked() else 0
                })
                self.db.insert_education({
                    'snils': self.inputSnils.text().strip(),
                    'school': self.inputSchool.text().strip(),
                    'graduation_year': self.inputGraduationYear.value(),
                    'average_grade': self.inputAverageGrade.value()
                })
                QMessageBox.information(
                    self, "Успех", "Абитуриент успешно добавлен!")

            self.accept()

        except Exception as e:
            QMessageBox.critical(
                self, "Ошибка", f"Ошибка при сохранении:\n{e}")


class ApplicationDialog(QDialog):
    def __init__(self, parent=None, db=None, application_id=None):
        super().__init__(parent)
        uic.loadUi('./ui/application_dialog.ui', self)

        self.inputSubmissionDate.setDate(QDate(2026, 6, 1))

        self.db = db
        self.application_id = application_id

        self.load_data()

        self.inputSpecialty.currentIndexChanged.connect(self.update_study_type)

        if self.application_id:
            self.setWindowTitle("Редактирование заявления")
            self.titleLabel.setText("Редактирование заявления")
            self.fill_data()

        self.btnSave.clicked.connect(self.save)
        self.btnCancel.clicked.connect(self.reject)

    def fill_data(self):
        with self.db.connection.cursor() as cursor:
            cursor.execute("""
                SELECT snils, speciality_code, submission_date, benefit
                FROM applications
                WHERE application_id = %s
            """, (self.application_id,))
            row = cursor.fetchone()

        if not row:
            return

        snils, speciality, date, benefit = row

        # установить абитуриента
        index = self.comboBox.findData(snils)
        if index >= 0:
            self.comboBox.setCurrentIndex(index)

        # установить специальность
        index = self.inputSpecialty.findData(speciality)
        if index >= 0:
            self.inputSpecialty.setCurrentIndex(index)

        # дата
        self.inputSubmissionDate.setDate(date)

        # льгота
        index = self.inputBenefit.findText(
            benefit if benefit else "Выберите льготу")
        if index >= 0:
            self.inputBenefit.setCurrentIndex(index)

    def load_data(self):
        self.load_applicants()
        self.load_specialties()

    def load_applicants(self):
        with self.db.connection.cursor() as cursor:
            cursor.execute("""
                SELECT snils, last_name, first_name, middle_name
                FROM applicants
            """)
            rows = cursor.fetchall()

        self.comboBox.clear()
        self.comboBox.addItem("Выберите абитуриента", None)

        for snils, last, first, middle in rows:
            fio = f"{last} {first} {middle or ''}".strip()
            text = f"{snils} - {fio}"

            self.comboBox.addItem(text, snils)

    def load_specialties(self):
        with self.db.connection.cursor() as cursor:
            cursor.execute("""
                SELECT speciality_code, speciality_name, budget_places, paid_places
                FROM specialties
            """)
            rows = cursor.fetchall()

        self.inputSpecialty.clear()
        self.inputSpecialty.addItem("Выберите специальность", None)

        self.specialty_data = {}

        for code, name, budget, paid in rows:
            text = f"{code} - {name}"
            self.inputSpecialty.addItem(text, code)

            # сохраняем места
            self.specialty_data[code] = {
                "budget": budget,
                "paid": paid
            }

    def save(self):
        # Валидация
        snils = self.comboBox.currentData()
        speciality = self.inputSpecialty.currentData()
        study_type = self.cmbStudyType.currentText().strip()
        benefit = self.inputBenefit.currentText()

        if self.db.application_exists(snils, speciality, study_type):
            QMessageBox.warning(
                self, "Ошибка", "Такое заявление уже существует")
            return

        if benefit == "Выберите льготу":
            benefit = None

        if not snils:
            QMessageBox.warning(self, "Ошибка", "Выберите абитуриента")
            return

        if not speciality:
            QMessageBox.warning(self, "Ошибка", "Выберите специальность")
            return

        if not study_type:
            QMessageBox.warning(self, "Ошибка", "Выберите тип обучения")
            return

        # Проверка мест
        data = self.specialty_data.get(speciality, {})
        budget = data.get("budget", 0)
        paid = data.get("paid", 0)

        if study_type == "Бюджет" and budget <= 0:
            QMessageBox.warning(self, "Ошибка", "Нет бюджетных мест")
            return

        if study_type == "Платное" and paid <= 0:
            QMessageBox.warning(self, "Ошибка", "Нет платных мест")
            return

        date = self.inputSubmissionDate.date().toString("yyyy-MM-dd")

        data = {
            'snils': snils,
            'speciality_code': speciality,
            'submission_date': date,
            'benefit': benefit,
            'study_type': study_type   # ← ВОТ ЭТОГО НЕ ХВАТАЛО
        }

        try:
            if self.application_id:
                self.db.update_application(self.application_id, data)
                QMessageBox.information(self, "Успех", "Заявление обновлено!")
            else:
                self.db.insert_application(data)
                QMessageBox.information(self, "Успех", "Заявление добавлено!")

            self.accept()

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка сохранения:\n{e}")

    def update_study_type(self):
        self.cmbStudyType.clear()

        code = self.inputSpecialty.currentData()
        if not code:
            return

        data = self.specialty_data.get(code, {})
        budget = data.get("budget", 0)
        paid = data.get("paid", 0)

        if budget > 0:
            self.cmbStudyType.addItem("Бюджет")

        if paid > 0:
            self.cmbStudyType.addItem("Платное")

        if budget == 0 and paid == 0:
            QMessageBox.warning(self, "Ошибка", "Нет доступных мест!")


class SpecialtiesDialog(QDialog):
    def __init__(self, parent=None, db=None, speciality_code=None):
        super().__init__(parent)
        uic.loadUi('./ui/direction_dialog.ui', self)

        self.db = db
        self.speciality_code = speciality_code

        self.load_data()

        if self.speciality_code:
            self.setWindowTitle("Редактирование направления")
            self.titleLabel.setText("Редактирование направления")
            self.fill_data()
        else:
            self.setWindowTitle("Добавление направления")

        self.btnSave.clicked.connect(self.save)
        self.btnCancel.clicked.connect(self.reject)

    def load_data(self):
        """Загрузка отделений в comboBox"""
        self.inputDepartment.clear()
        self.inputDepartment.addItem("Выберите отделение", None)

        with self.db.connection.cursor() as cursor:
            cursor.execute("""
                SELECT department_id, department_name
                FROM departaments
            """)
            rows = cursor.fetchall()

        for dep_id, dep_name in rows:
            self.inputDepartment.addItem(dep_name, dep_id)

    def fill_data(self):
        with self.db.connection.cursor() as cursor:
            cursor.execute("""
                SELECT speciality_code, speciality_name, department_id,
                    budget_places, paid_places, study_duration, study_form
                FROM specialties
                WHERE speciality_code = %s
            """, (self.speciality_code,))
            row = cursor.fetchone()

        if not row:
            return

        code, name, dep_id, budget, paid, duration, form = row

        self.inputSpecialityCode.setText(code)
        self.inputSpecialityName.setText(name)
        self.inputBudgetPlaces.setValue(budget)
        self.inputPaidPlaces.setValue(paid)
        self.lineEdit.setText(duration)

        # форма
        index = self.inputFormOfStudy.findText(form if form else "")
        if index >= 0:
            self.inputFormOfStudy.setCurrentIndex(index)

        # отделение
        index = self.inputDepartment.findData(dep_id)
        if index >= 0:
            self.inputDepartment.setCurrentIndex(index)

        # при редактировании код лучше запретить менять
        self.inputSpecialityCode.setEnabled(False)

    def validate(self):
        code = self.inputSpecialityCode.text().strip()
        name = self.inputSpecialityName.text().strip()
        dep_id = self.inputDepartment.currentData()
        duration = self.lineEdit.text().strip()
        form = self.inputFormOfStudy.currentText()
        duration = self.lineEdit.text().strip()

        parts = duration.split()

        # 1 год 10 мес
        if len(parts) != 4:
            QMessageBox.warning(
                self,
                "Ошибка",
                "Срок обучения: формат '1 год 10 мес'"
            )
            return False

        num1, word1, num2, word2 = parts

        # проверка чисел
        if not num1.isdigit() or not num2.isdigit():
            QMessageBox.warning(
                self,
                "Ошибка",
                "Годы и месяцы должны быть числами"
            )
            return False

        num1 = int(num1)
        num2 = int(num2)

        # логические ограничения
        if num1 < 0 or num1 > 6:
            QMessageBox.warning(self, "Ошибка", "Годы должны быть от 0 до 6")
            return False

        if num2 < 0 or num2 > 11:
            QMessageBox.warning(
                self, "Ошибка", "Месяцы должны быть от 0 до 11")
            return False

        # проверка слов
        valid_years = ["год", "года", "лет"]
        valid_months = ["мес", "месяц", "месяца"]

        if word1 not in valid_years:
            QMessageBox.warning(self, "Ошибка", "Неверное слово для лет")
            return False

        if word2 not in valid_months:
            QMessageBox.warning(self, "Ошибка", "Неверное слово для месяцев")
            return False

        if not code:
            QMessageBox.warning(self, "Ошибка", "Введите код специальности")
            return False

        if not name:
            QMessageBox.warning(
                self, "Ошибка", "Введите название специальности")
            return False

        if not dep_id:
            QMessageBox.warning(self, "Ошибка", "Выберите отделение")
            return False

        # код 00.00.00
        parts = code.split('.')
        if len(parts) != 3 or not all(p.isdigit() for p in parts):
            QMessageBox.warning(
                self, "Ошибка", "Код должен быть в формате 00.00.00")
            return False

        # форма обучения
        if form not in ["Очная", "Заочная"]:
            QMessageBox.warning(self, "Ошибка", "Некорректная форма обучения")
            return False

        return True

    def save(self):
        if not self.validate():
            return

        code = self.inputSpecialityCode.text().strip()
        name = self.inputSpecialityName.text().strip()
        dep_id = self.inputDepartment.currentData()
        budget = self.inputBudgetPlaces.value()
        paid = self.inputPaidPlaces.value()
        duration = self.lineEdit.text().strip()
        form = self.inputFormOfStudy.currentText()

        data = {
            'code': code,
            'name': name,
            'dep_id': dep_id,
            'budget': budget,
            'paid': paid,
            'duration': duration,
            'form': form
        }

        try:
            if self.speciality_code:
                self.db.update_specialty(self.speciality_code, data)
            else:
                self.db.insert_specialty(data)

            QMessageBox.information(self, "Успех", "Сохранено!")
            self.accept()

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", str(e))


class DepartmentDialog(QDialog):
    """Диалог добавления/редактирования отделений"""

    def __init__(self, parent=None, db=None, department_id=None):
        super().__init__(parent)
        uic.loadUi('./ui/departments_dialog.ui', self)

        self.db = db
        self.department_id = department_id

        if self.department_id:
            self.setWindowTitle("Редактирование отделения")
            self.titleLabel.setText("Редактирование отделения")
            self.fill_data()
        else:
            self.setWindowTitle("Добавление отделения")
            self.titleLabel.setText("Добавление отделения")

        self.btnSave.clicked.connect(self.save)
        self.btnCancel.clicked.connect(self.reject)

    def fill_data(self):
        """Заполнение полей данными из БД"""
        with self.db.connection.cursor() as cursor:
            cursor.execute(
                "SELECT department_id, department_name, head_of_department FROM departaments WHERE department_id = %s",
                (self.department_id,)
            )
            row = cursor.fetchone()

        if not row:
            return

        department_id, name, head = row

        self.inputDepartmentName.setText(str(name))
        self.inputDepartmentHead.setText(str(head) if head else '')

    def validate(self):
        """Валидация данных"""
        name = self.inputDepartmentName.text().strip()
        head = self.inputDepartmentHead.text().strip()

        if not name:
            QMessageBox.warning(self, "Ошибка", "Введите название отделения")
            return False

        if not head:
            QMessageBox.warning(self, "Ошибка", "Введите ФИО заведующего")
            return False

        return True

    def save(self):
        """Сохранение данных"""
        if not self.validate():
            return

        name = self.inputDepartmentName.text().strip()
        head = self.inputDepartmentHead.text().strip()

        data = {
            'name': name,
            'head': head
        }

        try:
            if self.department_id:
                self.db.update_department(self.department_id, data)
                QMessageBox.information(self, "Успех", "Отделение обновлено!")
            else:
                self.db.insert_department(data)
                QMessageBox.information(self, "Успех", "Отделение добавлено!")

            self.accept()

        except Exception as e:
            QMessageBox.critical(
                self, "Ошибка", f"Ошибка при сохранении:\n{e}")


class MyWidget(QMainWindow):
    def __init__(self, role):
        super().__init__()
        uic.loadUi('./ui/main_window.ui', self)

        self.db = Database()

        self.setup_tables()
        self.role = role

        self.search_config = {
            'applicants': {
                'search_field': self.searchApplicants,
                'table_widget': self.tableApplicants
            },
            'applications': {
                'search_field': self.searchApplication,
                'table_widget': self.tableApplications
            },
            'departaments': {
                'search_field': self.searchDepartments,
                'table_widget': self.tableDepartments
            },
            'specialties': {
                'search_field': self.searchDirections,
                'table_widget': self.tableDirections
            }
        }

        self.connect_button()
        self.load_combox()
        self.load_data()
        self.setup_permissions()

        self.menu_buttons = [
            self.btnApplicants,
            self.btnApplications,
            self.btnDirections,
            self.btnDepartments,
            self.btnReports
        ]
        self.btnApplicants.setChecked(True)
        self.stackedWidget.setCurrentWidget(self.pageApplicants)
        # self.btnReports.clicked.connect(self.run)

    def on_menu_click(self, btn):
        for b in self.menu_buttons:
            b.setChecked(False)  # снять со всех
        btn.setChecked(True)  # поставить на нажатую

    def connect_button(self):
        '''Подключение кнопок'''
        self.btnApplicants.clicked.connect(
            lambda: (self.stackedWidget.setCurrentWidget(self.pageApplicants), self.on_menu_click(self.btnApplicants)))
        self.btnApplications.clicked.connect(
            lambda: (self.stackedWidget.setCurrentWidget(self.pageApplications), self.on_menu_click(self.btnApplications)))
        self.btnDirections.clicked.connect(
            lambda: (self.stackedWidget.setCurrentWidget(self.pageDirections), self.on_menu_click(self.btnDirections)))
        self.btnDepartments.clicked.connect(
            lambda: (self.stackedWidget.setCurrentWidget(self.pageDepartments), self.on_menu_click(self.btnDepartments)))
        self.btnReports.clicked.connect(
            lambda: (self.stackedWidget.setCurrentWidget(self.pageReports), self.on_menu_click(self.btnReports)))
        self.btnSearchApplicant.clicked.connect(lambda:
                                                (self.search_and_filter('applicants')))
        self.btnSearchApplicataion.clicked.connect(lambda:
                                                   self.search_and_filter('applications'))
        self.btnSearchDirection.clicked.connect(lambda:
                                                self.search_and_filter('specialties'))
        self.btnSearchDepartment.clicked.connect(lambda:
                                                 self.search_and_filter('departaments'))
        self.btnUpdateApplicant.clicked.connect(
            lambda: (self.load_data(), self.reset_selection(self.tableApplicants, self.searchApplicants), self.cmbGender.setCurrentIndex(0), self.cmbCertificate.setCurrentIndex(0), self.cmbPhoto.setCurrentIndex(0)))

        self.btnUpdateApplication.clicked.connect(
            lambda: (self.load_data(), self.reset_selection(self.tableApplications, self.searchApplication), self.cmbBenefit.setCurrentIndex(0)))

        self.btnUpdateDirection.clicked.connect(
            lambda: (self.load_data(), self.reset_selection(self.tableDirections, self.searchDirections)))

        self.btnUpdateDepartment.clicked.connect(
            lambda: (self.load_data(), self.reset_selection(self.tableDepartments, self.searchDepartments), ))
        self.btnAddApplicant.clicked.connect(
            lambda: self.open_add_dialog(ApplicantDialog))
        self.btnAddApplication.clicked.connect(
            lambda: self.open_add_dialog(ApplicationDialog))
        self.btnAddDirection.clicked.connect(
            lambda: self.open_add_dialog(SpecialtiesDialog))
        self.btnAddDepartment.clicked.connect(
            lambda: self.open_add_dialog(DepartmentDialog))
        self.btnEditApplicant.clicked.connect(self.open_edit_dialog)
        self.btnEditApplication.clicked.connect(self.open_edit_dialog)
        self.btnEditDirection.clicked.connect(self.open_edit_dialog)
        self.btnEditDepartment.clicked.connect(self.open_edit_dialog)
        self.btnDeleteApplicant.clicked.connect(self.delete_applicant)
        self.btnDeleteApplication.clicked.connect(self.delete_application)
        self.btnDeleteDirection.clicked.connect(self.delete_direction)
        self.btnDeleteDepartment.clicked.connect(self.delete_department)
        self.btnRepDirections.clicked.connect(self.generate_competition_doc)
        self.btnRepEnrollment.clicked.connect(
            self.generate_enrollment_split_tables)

    def setup_permissions(self):
        if self.role == "operator":
            # ❌ отчёты
            self.btnReports.setEnabled(False)

            # ❌ направления
            self.btnAddDirection.setEnabled(False)
            self.btnEditDirection.setEnabled(False)
            self.btnDeleteDirection.setEnabled(False)

            # ❌ кнопки отделений (если есть)
            if hasattr(self, 'btnAddDepartment'):
                self.btnAddDepartment.setEnabled(False)
            if hasattr(self, 'btnEditDepartment'):
                self.btnEditDepartment.setEnabled(False)
            if hasattr(self, 'btnDeleteDepartment'):
                self.btnDeleteDepartment.setEnabled(False)

    def setup_tables(self):
        '''Загрузка заголовков таблиц'''
        self.tableApplicants.setColumnCount(12)
        self.tableApplicants.setHorizontalHeaderLabels(
            ['СНИЛС', 'Фамилия', 'Имя', 'Отчество', 'Дата рождения', 'Пол', 'Номер телефона', 'Паспортные данные', 'Медицинская справка', 'Электронная почта', 'Адрес', 'Фото'])
        self.tableApplicants.horizontalHeader().setSectionResizeMode(
            QHeaderView.ResizeMode.ResizeToContents)

        self.tableApplications.setColumnCount(6)
        self.tableApplications.setHorizontalHeaderLabels(
            ['id', 'СНИЛС', 'Название специальности', 'Дата подачи', 'Льгота', 'Тип обучения'])
        self.tableApplications.horizontalHeader().setSectionResizeMode(
            QHeaderView.ResizeMode.ResizeToContents)
        self.tableApplications.setColumnHidden(0, True)

        self.tableDirections.setColumnCount(7)
        self.tableDirections.setHorizontalHeaderLabels(
            ['Код специальности', 'Название специальности', 'Название отделения', 'Количество бюджетных мест', 'Количество платных мест', 'Время обучения', 'Форма обучения'])
        self.tableDirections.horizontalHeader().setSectionResizeMode(
            QHeaderView.ResizeMode.ResizeToContents)
        # self.tableDirections.setColumnHidden(0, True)

        self.tableDepartments.setColumnCount(3)
        self.tableDepartments.setHorizontalHeaderLabels(
            ['id', 'Название отделения', 'Заведующий отделения'])
        self.tableDepartments.horizontalHeader().setSectionResizeMode(
            QHeaderView.ResizeMode.ResizeToContents)
        self.tableDepartments.setColumnHidden(0, True)

    def format_value(self, data, column_number):
        """Форматирование значений для определённых столбцов"""
        # Номера столбцов, которые нужно форматировать
        format_columns = {8, 11}
        print(data, type(data))
        if data == None:
            return '-'
        if column_number in format_columns:
            if str(data) == '0':
                return 'Нет'
            elif str(data) == '1':
                return 'Есть'
        return data

    def load_data(self):
        '''Загрузка данных из бд'''

        tables = {
            'applicants': self.tableApplicants,
            'applications': self.tableApplications,
            'departaments': self.tableDepartments,
            'specialties': self.tableDirections
        }

        for key, value in tables.items():
            print(key, value)
            if key == 'applications':
                sql = self.db.get_applications_with_specialties()
            elif key == 'specialties':
                sql = self.db.get_specialties_with_departaments()
            else:
                sql = self.db.select(key)

            if sql and len(sql) > 0:
                actual_columns = len(sql[0])
                print(
                    f"База вернула {actual_columns} колонок, таблица ожидает {value.columnCount()}")

                if actual_columns != value.columnCount():
                    value.setColumnCount(actual_columns)

            value.setRowCount(len(sql))
            for row_number, row_data in enumerate(sql):
                for column_number, data in enumerate(row_data):
                    formatted = self.format_value(data, column_number)
                    value.setItem(row_number, column_number,
                                  QTableWidgetItem(str(formatted)))

    def load_combox(self):
        with self.db.connection.cursor() as cursor:
            cursor.execute("""
                SELECT department_name
                FROM departaments
            """)
            rows = cursor.fetchall()

        for department_name in rows:
            name = department_name
            self.cmbDirection.addItem(str(*name))

    def search_and_filter(self, table_name):
        config = self.search_config[table_name]
        table_widget = config['table_widget']
        search_text = config['search_field'].text().lower()

        if table_name == 'applicants':
            filters = {
                'gender': self.cmbGender.currentText(),
                'medical_certificate': '1' if self.cmbCertificate.currentText() == 'Есть' else
                ('0' if self.cmbCertificate.currentText()
                 == 'Нет' else 'Выберите фото'),
                'foto': '1' if self.cmbPhoto.currentText() == 'Есть' else
                ('0' if self.cmbPhoto.currentText() == 'Нет' else 'Выберите фото')
            }
            all_default = all('Выберите' in v for v in filters.values())
            if all_default:
                result = self.db.select(table_name)
            else:
                result = self.db.filter(table_name, filters)

        elif table_name == 'applications':
            filters = {'benefit': self.cmbBenefit.currentText()}
            all_default = all('Выберите' in v for v in filters.values())
            if all_default:
                result = self.db.get_applications_with_specialties()
            else:
                result = self.db.filter(table_name, filters)

        elif table_name == 'specialties':
            filters = {'department_name': self.cmbDirection.currentText()}
            all_default = all('Выберите' in v for v in filters.values())
            if all_default:
                result = self.db.get_specialties_with_departaments()
            else:
                result = self.db.filter(table_name, filters)

        else:
            result = self.db.select(table_name)

        table_widget.setRowCount(len(result))
        for row_number, row_data in enumerate(result):
            for column_number, data in enumerate(row_data):
                formatted = self.format_value(data, column_number)
                table_widget.setItem(row_number, column_number,
                                     QTableWidgetItem(str(formatted)))

        if search_text:
            table_widget.clearSelection()
            table_widget.setSelectionMode(
                table_widget.SelectionMode.MultiSelection
            )
            for row_number, row_data in enumerate(result):
                for data in row_data:
                    if search_text in str(data).lower():
                        table_widget.selectRow(row_number)
                        break
        else:
            table_widget.clearSelection()
            table_widget.setSelectionMode(
                table_widget.SelectionMode.SingleSelection
            )

    def reset_selection(self, table_widget, search_field):
        table_widget.clearSelection()
        table_widget.setSelectionMode(
            table_widget.SelectionMode.SingleSelection)
        search_field.clear()

    def open_add_dialog(self, dialog_class):
        dialog = dialog_class(parent=self, db=self.db)
        if dialog.exec():
            self.load_data()

    def open_edit_dialog(self):
        current_page = self.stackedWidget.currentWidget()

        # Абитуриенты
        if current_page == self.pageApplicants:
            selected = self.tableApplicants.selectedItems()
            if not selected:
                QMessageBox.warning(self, "Ошибка", "Выберите абитуриента")
                return

            row = self.tableApplicants.currentRow()
            snils = self.tableApplicants.item(row, 0).text()

            dialog = ApplicantDialog(self.db, snils=snils, parent=self)
            if dialog.exec():
                self.load_data()

        # Заявления
        elif current_page == self.pageApplications:
            selected = self.tableApplications.selectedItems()
            if not selected:
                QMessageBox.warning(self, "Ошибка", "Выберите заявление")
                return

            row = self.tableApplications.currentRow()
            app_id = int(self.tableApplications.item(row, 0).text())
            snils = self.tableApplications.item(row, 1).text()
            speciality_name = self.tableApplications.item(row, 2).text()
            date = self.tableApplications.item(row, 3).text()
            benefit = self.tableApplications.item(row, 4).text()

            # ← передаём application_id сразу, чтобы диалог знал что это редактирование
            dialog = ApplicationDialog(self, db=self.db, application_id=app_id)

            # СНИЛС
            index = dialog.comboBox.findData(snils)
            if index >= 0:
                dialog.comboBox.setCurrentIndex(index)

            # специальность
            for i in range(dialog.inputSpecialty.count()):
                if speciality_name in dialog.inputSpecialty.itemText(i):
                    dialog.inputSpecialty.setCurrentIndex(i)
                    break

            # дата
            dialog.inputSubmissionDate.setDate(
                QDate.fromString(date, "yyyy-MM-dd")
            )

            # льгота
            if benefit and benefit != '-':
                index = dialog.inputBenefit.findText(benefit)
                if index >= 0:
                    dialog.inputBenefit.setCurrentIndex(index)

            if dialog.exec():
                self.load_data()

        # Специальность
        elif current_page == self.pageDirections:
            selected = self.tableDirections.selectedItems()
            if not selected:
                QMessageBox.warning(self, "Ошибка", "Выберите направление")
                return

            row = self.tableDirections.currentRow()

            # берём код специальности
            speciality_code = self.tableDirections.item(row, 0).text()

            dialog = SpecialtiesDialog(
                parent=self,
                db=self.db,
                speciality_code=speciality_code
            )

            if dialog.exec():
                self.load_data()

        # Отделение
        elif current_page == self.pageDepartments:
            selected = self.tableDepartments.selectedItems()
            if not selected:
                QMessageBox.warning(self, "Ошибка", "Выберите отделение")
                return

            row = self.tableDepartments.currentRow()

            # берём ID отделения из скрытого столбца
            department_id = int(self.tableDepartments.item(row, 0).text())

            dialog = DepartmentDialog(
                parent=self,
                db=self.db,
                department_id=department_id
            )

            if dialog.exec():
                self.load_data()

    def delete_applicant(self):
        '''Удаление абитуриента'''
        selected = self.tableApplicants.selectedItems()
        if not selected:
            QMessageBox.warning(
                self, "Ошибка", "Выберите абитуриента для удаления")
            return

        row = self.tableApplicants.currentRow()
        snils = self.tableApplicants.item(row, 0).text()

        reply = QMessageBox.question(
            self, "Подтверждение",
            f"Вы уверены что хотите удалить абитуриента?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            try:
                self.db.delete_applicant(snils)
                QMessageBox.information(
                    self, "Успех", "Абитуриент успешно удалён!")
                self.load_data()
            except Exception as e:
                QMessageBox.critical(
                    self, "Ошибка", f"Ошибка при удалении:\n{e}")

    def delete_application(self):
        '''Удаление заявления'''
        selected = self.tableApplications.selectedItems()
        if not selected:
            QMessageBox.warning(
                self, "Ошибка", "Выберите заявление для удаления")
            return

        row = self.tableApplications.currentRow()
        # ID заявления из скрытого столбца
        app_id = int(self.tableApplications.item(row, 0).text())

        reply = QMessageBox.question(
            self, "Подтверждение",
            "Вы уверены, что хотите удалить это заявление?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            try:
                self.db.delete_application(app_id)
                QMessageBox.information(
                    self, "Успех", "Заявление успешно удалено!")
                self.load_data()  # Обновляем таблицу
            except Exception as e:
                QMessageBox.critical(
                    self, "Ошибка", f"Ошибка при удалении:\n{e}")

    def delete_direction(self):
        selected = self.tableDirections.selectedItems()
        if not selected:
            QMessageBox.warning(
                self, "Ошибка", "Выберите направление для удаления")
            return

        row = self.tableDirections.currentRow()
        speciality_code = self.tableDirections.item(row, 0).text()

        reply = QMessageBox.question(
            self,
            "Подтверждение",
            "Вы уверены, что хотите удалить это направление?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            try:
                with self.db.connection.cursor() as cursor:
                    cursor.execute(
                        "DELETE FROM specialties WHERE speciality_code = %s",
                        (speciality_code,)
                    )
                self.db.connection.commit()

                QMessageBox.information(self, "Успех", "Направление удалено!")
                self.load_data()

            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Ошибка:\n{e}")

    def delete_department(self):
        """Удаление отделения"""
        selected = self.tableDepartments.selectedItems()
        if not selected:
            QMessageBox.warning(
                self, "Ошибка", "Выберите отделение для удаления")
            return

        row = self.tableDepartments.currentRow()
        department_id = int(self.tableDepartments.item(row, 0).text())

        reply = QMessageBox.question(
            self,
            "Подтверждение",
            "Вы уверены, что хотите удалить это отделение?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            try:
                self.db.delete_department(department_id)
                QMessageBox.information(self, "Успех", "Отделение удалено!")
                self.load_data()

            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Ошибка:\n{e}")

    def generate_competition_doc(self):
        data = self.db.get_competition_list()

        doc = Document()
        doc.add_heading("Конкурсный список абитуриентов", 0)

        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'

        headers = [
            "СНИЛС", "ФИО", "Специальность",
            "Тип обучения", "Льгота", "Дата подачи", "Средний балл"
        ]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for row in data:
            snils, last, first, middle, spec, study, benefit, date, grade = row
            fio = f"{last} {first} {middle or ''}".strip()
            cells = table.add_row().cells
            cells[0].text = str(snils)
            cells[1].text = fio
            cells[2].text = spec
            cells[3].text = study or "-"
            cells[4].text = benefit if benefit else "-"
            cells[5].text = str(date)
            cells[6].text = str(grade)

        doc.save("competition_list.docx")
        QMessageBox.information(self, "Успех", "Конкурсный список создан!")

    def generate_enrollment_split_tables(self):
        data = self.db.get_enrollment_data()

        # Группировка по специальностям
        specialties = {}
        for row in data:
            snils, last, first, middle, spec, study, benefit, date, grade, budget_places, paid_places = row
            if spec not in specialties:
                specialties[spec] = {
                    "budget_limit": budget_places,
                    "paid_limit":   paid_places,
                    "applicants":   []
                }
            specialties[spec]["applicants"].append(row)

        doc = Document()
        doc.add_heading("ПРИКАЗ О ЗАЧИСЛЕНИИ", 0)

        for spec, info in specialties.items():
            budget_limit = info["budget_limit"]
            paid_limit = info["paid_limit"]

            applicants = sorted(
                info["applicants"],
                key=lambda x: (
                    0 if x[6] else 1,     # льготники вверх
                    -float(x[8])          # был -x[8], падало на Decimal
                )
            )

            budget_list = []
            paid_list = []
            budget_count = 0
            paid_count = 0

            for row in applicants:
                snils, last, first, middle, _spec, study, benefit, date, grade, _bp, _pp = row
                fio = f"{last} {first} {middle or ''}".strip()

                if benefit and budget_count < budget_limit:
                    budget_list.append((snils, fio, float(grade), benefit))
                    budget_count += 1
                elif study == "Бюджет" and budget_count < budget_limit:
                    budget_list.append((snils, fio, float(grade), benefit))
                    budget_count += 1
                elif study == "Платное" and paid_count < paid_limit:
                    paid_list.append((snils, fio, float(grade), benefit))
                    paid_count += 1

            if not budget_list and not paid_list:
                continue

            doc.add_heading(f"Специальность: {spec}", level=1)

            headers = ["СНИЛС", "ФИО", "Средний балл", "Льгота"]

            doc.add_heading("Бюджет", level=2)
            if budget_list:
                table_b = doc.add_table(rows=1, cols=4)
                table_b.style = 'Table Grid'
                for i, h in enumerate(headers):
                    table_b.rows[0].cells[i].text = h
                for snils, fio, grade, benefit in budget_list:
                    cells = table_b.add_row().cells
                    cells[0].text = str(snils)
                    cells[1].text = fio
                    cells[2].text = f"{grade:.2f}"
                    cells[3].text = benefit if benefit else "-"
            else:
                doc.add_paragraph("Нет зачисленных на бюджет.")

            # --- ПЛАТНО ---
            doc.add_heading("Платное обучение", level=2)
            if paid_list:
                table_p = doc.add_table(rows=1, cols=4)
                table_p.style = 'Table Grid'
                for i, h in enumerate(headers):
                    table_p.rows[0].cells[i].text = h
                for snils, fio, grade, benefit in paid_list:
                    cells = table_p.add_row().cells
                    cells[0].text = str(snils)
                    cells[1].text = fio
                    cells[2].text = f"{grade:.2f}"
                    cells[3].text = benefit if benefit else "-"
            else:
                doc.add_paragraph("Нет зачисленных на платное обучение.")

            doc.add_paragraph("")  # отступ между специальностями

        doc.save("prikaz_o_zachislenii.docx")
        QMessageBox.information(self, "Готово", "Приказ сформирован успешно!")


if __name__ == '__main__':
    # db = Database()
    # db.select()
    # db.close()

    app = QApplication(sys.argv)
    ex = LoginWindow()
    ex.show()
    sys.exit(app.exec())
